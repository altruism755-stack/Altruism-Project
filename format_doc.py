import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document(r'C:\Users\vvmar\Downloads\Graduation_Project_Filled (6).docx')

# ─── COLOURS ──────────────────────────────────────────────────────────────────
C_NAVY  = RGBColor(0x1F, 0x38, 0x64)   # chapter headings
C_DARK  = RGBColor(0x1A, 0x1A, 0x2E)   # section headings
C_BLACK = RGBColor(0x00, 0x00, 0x00)   # body text

BORDER_COLOR = '1F3864'  # same navy for page borders

# ─── XML HELPERS ──────────────────────────────────────────────────────────────

def get_or_add(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el

def set_spacing(pPr, line='360', line_rule='auto', before='0', after='120'):
    sp = get_or_add(pPr, 'w:spacing')
    sp.set(qn('w:line'), line)
    sp.set(qn('w:lineRule'), line_rule)
    sp.set(qn('w:before'), before)
    sp.set(qn('w:after'), after)

def set_indent(pPr, first_line='720', left=None):
    ind = get_or_add(pPr, 'w:ind')
    ind.set(qn('w:firstLine'), first_line)
    if left is not None:
        ind.set(qn('w:left'), left)
    elif qn('w:left') in ind.attrib:
        del ind.attrib[qn('w:left')]

def clear_first_line_indent(pPr):
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    if qn('w:firstLineChars') in ind.attrib:
        del ind.attrib[qn('w:firstLineChars')]

def set_jc(pPr, val='both'):
    jc = get_or_add(pPr, 'w:jc')
    jc.set(qn('w:val'), val)

def set_run_font(run, name='Times New Roman', size_pt=12, bold=None, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # Also set rFonts via XML for full compatibility
    r = run._r
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

# ═══════════════════════════════════════════════════════════════════════════════
#  1. UPDATE DOCUMENT-WIDE STYLES
# ═══════════════════════════════════════════════════════════════════════════════

def configure_style(style_name, font_pt, bold, color, line='360',
                    sp_before='0', sp_after='120',
                    align='both', first_line='720',
                    page_break_before=False, keep_with_next=False):
    if style_name not in [s.name for s in doc.styles]:
        return
    st = doc.styles[style_name]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(font_pt)
    st.font.bold = bold
    st.font.color.rgb = color

    pPr = st.element.get_or_add_pPr()

    # Alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), align)

    # Spacing
    set_spacing(pPr, line=line, before=sp_before, after=sp_after)

    # First-line indent
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if first_line == '0':
        ind.set(qn('w:firstLine'), '0')
    else:
        ind.set(qn('w:firstLine'), first_line)

    # Page break before
    pbr = pPr.find(qn('w:pageBreakBefore'))
    if page_break_before:
        if pbr is None:
            pbr = OxmlElement('w:pageBreakBefore')
            pPr.append(pbr)
    else:
        if pbr is not None:
            pPr.remove(pbr)

    # Keep with next
    kwn = pPr.find(qn('w:keepNext'))
    if keep_with_next:
        if kwn is None:
            kwn = OxmlElement('w:keepNext')
            pPr.append(kwn)

    # Also set font via rPr in style element for run-level defaults
    rPr = st.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        st.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

# Normal: 12pt, justified, 1.5 line, 0.5" first-line indent (720 twips)
configure_style('Normal', 12, False, C_BLACK, line='360',
                sp_before='0', sp_after='120', align='both', first_line='720')

# Normal (Web): same as Normal
configure_style('Normal (Web)', 12, False, C_BLACK, line='360',
                sp_before='0', sp_after='120', align='both', first_line='720')

# Heading 1: 14pt Bold, navy, left-aligned, page break before
configure_style('Heading 1', 14, True, C_NAVY, line='360',
                sp_before='480', sp_after='240', align='left', first_line='0',
                page_break_before=True, keep_with_next=True)

# Heading 2: 13pt Bold, dark, left-aligned
configure_style('Heading 2', 13, True, C_DARK, line='360',
                sp_before='240', sp_after='120', align='left', first_line='0',
                keep_with_next=True)

# Heading 3: 12pt Bold, dark, left-aligned
configure_style('Heading 3', 12, True, C_DARK, line='360',
                sp_before='160', sp_after='80', align='left', first_line='0',
                keep_with_next=True)

# List Paragraph: 12pt, justified, 1.5 line, indented left 0.5"
configure_style('List Paragraph', 12, False, C_BLACK, line='360',
                sp_before='0', sp_after='80', align='both', first_line='0')
if 'List Paragraph' in [s.name for s in doc.styles]:
    lp = doc.styles['List Paragraph']
    lp.paragraph_format.left_indent = Inches(0.5)

print("OK 1: Document styles configured")

# ═══════════════════════════════════════════════════════════════════════════════
#  2. PAGE LAYOUT: A4 + MARGINS FOR ALL SECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

for section in doc.sections:
    section.page_width        = Inches(8.27)
    section.page_height       = Inches(11.69)
    section.top_margin        = Inches(1.0)
    section.bottom_margin     = Inches(1.0)
    section.left_margin       = Inches(1.25)
    section.right_margin      = Inches(1.0)
    section.header_distance   = Inches(0.5)
    section.footer_distance   = Inches(0.5)

print("OK 2: A4 page layout set for all sections")

# ═══════════════════════════════════════════════════════════════════════════════
#  3. ELEGANT PAGE BORDERS (single-line, navy, all pages)
# ═══════════════════════════════════════════════════════════════════════════════

def add_page_borders(sectPr):
    for old in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(old)
    pgB = OxmlElement('w:pgBorders')
    pgB.set(qn('w:offsetFrom'), 'page')
    for edge in ('top', 'left', 'bottom', 'right'):
        bd = OxmlElement(f'w:{edge}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), '12')      # 1.5 pt line (sz in 1/8pt units)
        bd.set(qn('w:space'), '24')   # 24pt from page edge
        bd.set(qn('w:color'), BORDER_COLOR)
        pgB.append(bd)
    # Insert before w:pgSz
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.addprevious(pgB)
    else:
        sectPr.append(pgB)

for section in doc.sections:
    add_page_borders(section._sectPr)

print("OK 3: Page borders added to all sections")

# ═══════════════════════════════════════════════════════════════════════════════
#  4. PAGE NUMBERING: Roman for preliminary, Arabic from Chapter 1
# ═══════════════════════════════════════════════════════════════════════════════

def set_pg_num_type(sectPr, fmt, start=None):
    for old in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(old)
    el = OxmlElement('w:pgNumType')
    el.set(qn('w:fmt'), fmt)
    if start is not None:
        el.set(qn('w:start'), str(start))
    sectPr.append(el)

sections = doc.sections
# Section 0 → cover page: roman starting at i (no footer shown)
set_pg_num_type(sections[0]._sectPr, 'lowerRoman', start=1)
# Section 1 → preliminary: roman starting at ii (Declaration = ii per TOC)
if len(sections) > 1:
    set_pg_num_type(sections[1]._sectPr, 'lowerRoman', start=2)
# Section 2 → main chapters: arabic starting at 1
if len(sections) > 2:
    set_pg_num_type(sections[2]._sectPr, 'decimal', start=1)

print("OK 4: Page numbering configured (roman prelim / arabic chapters)")

# ═══════════════════════════════════════════════════════════════════════════════
#  5. FOOTER WITH CENTERED PAGE NUMBERS
# ═══════════════════════════════════════════════════════════════════════════════

def add_page_num_footer(section):
    footer = section.footer

    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr = fp._p.get_or_add_pPr()
    set_spacing(pPr, line='240', before='100', after='0')  # single in footer
    clear_first_line_indent(pPr)

    run = fp.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    # Insert PAGE field
    for tag, attr, text in [
        ('w:fldChar', {qn('w:fldCharType'): 'begin'}, None),
        ('w:instrText', {'{http://www.w3.org/XML/1998/namespace}space': 'preserve'}, ' PAGE '),
        ('w:fldChar', {qn('w:fldCharType'): 'end'}, None),
    ]:
        el = OxmlElement(tag)
        for k, v in attr.items():
            el.set(k, v)
        if text:
            el.text = text
        run._r.append(el)

for section in doc.sections:
    add_page_num_footer(section)

print("OK 5: Page number footers added to all sections")

# ═══════════════════════════════════════════════════════════════════════════════
#  6. FIX COVER PAGE PARAGRAPHS (centered, no indent, no first-line indent)
# ═══════════════════════════════════════════════════════════════════════════════

# Cover page = paragraphs before first real Heading 1 ("Declaration") at index 30
for i in range(30):
    p = doc.paragraphs[i]
    pPr = p._p.get_or_add_pPr()
    set_jc(pPr, 'center')
    clear_first_line_indent(pPr)
    set_spacing(pPr, line='360', before='0', after='120')
    for run in p.runs:
        if not run.font.size:
            run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

print("OK 6: Cover page centered and formatted")

# ═══════════════════════════════════════════════════════════════════════════════
#  7. FIX HEADING 1 RUNS: font + page-break-before + no indent
# ═══════════════════════════════════════════════════════════════════════════════

# "Date: …… / .….. / 2026" is incorrectly styled as Heading 1 — fix it
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and p.text.strip().startswith('Date:'):
        p.style = doc.styles['Normal']
        pPr = p._p.get_or_add_pPr()
        clear_first_line_indent(pPr)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = False

for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        pPr = p._p.get_or_add_pPr()
        # Ensure page break before
        if pPr.find(qn('w:pageBreakBefore')) is None:
            pbr = OxmlElement('w:pageBreakBefore')
            pPr.append(pbr)
        # No indent
        clear_first_line_indent(pPr)
        ind = pPr.find(qn('w:ind'))
        if ind is not None and qn('w:left') in ind.attrib:
            del ind.attrib[qn('w:left')]
        # Fix runs
        for run in p.runs:
            set_run_font(run, 'Times New Roman', 14, True, C_NAVY)

print("OK 7: Heading 1 paragraphs formatted")

# ═══════════════════════════════════════════════════════════════════════════════
#  8. FIX HEADING 2 AND 3 RUNS
# ═══════════════════════════════════════════════════════════════════════════════

for p in doc.paragraphs:
    if p.style.name == 'Heading 2':
        pPr = p._p.get_or_add_pPr()
        clear_first_line_indent(pPr)
        for run in p.runs:
            set_run_font(run, 'Times New Roman', 13, True, C_DARK)
    elif p.style.name == 'Heading 3':
        pPr = p._p.get_or_add_pPr()
        clear_first_line_indent(pPr)
        for run in p.runs:
            set_run_font(run, 'Times New Roman', 12, True, C_DARK)

print("OK 8: Heading 2 and 3 formatted")

# ═══════════════════════════════════════════════════════════════════════════════
#  9. FIX TOC + LIST-OF ENTRIES: remove first-line indent, light spacing
# ═══════════════════════════════════════════════════════════════════════════════

# These ranges are TOC entries and preliminary list entries that must not indent
# Identified from inspection:
# 55-105  → Table of Contents entries
# 107-113 → List of Tables entries
# 116-124 → List of Figures entries
# 126-127 → List of Abbreviations entries (short)

toc_ranges = list(range(55, 106)) + list(range(107, 114)) + list(range(116, 125)) + list(range(126, 130))

for i in toc_ranges:
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    pPr = p._p.get_or_add_pPr()
    clear_first_line_indent(pPr)
    set_spacing(pPr, line='300', before='0', after='60')
    set_jc(pPr, 'both')
    for run in p.runs:
        if not run.font.size:
            run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

print("OK 9: TOC/list entries formatted (no first-line indent)")

# ═══════════════════════════════════════════════════════════════════════════════
#  10. FIX DECLARATION / ACKNOWLEDGEMENT / ABSTRACT BODY PARAGRAPHS
#      These should have NO first-line indent (they are centered or special)
# ═══════════════════════════════════════════════════════════════════════════════

# Declaration body: paragraphs 31..41 (after "Declaration" H1, before "Acknowledgements" H1)
# These include signatures, supervisor approval — should be no indent or left aligned
declaration_body = list(range(31, 42))
for i in declaration_body:
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
        continue
    pPr = p._p.get_or_add_pPr()
    clear_first_line_indent(pPr)
    set_jc(pPr, 'both')

print("OK 10: Declaration body formatted")

# ═══════════════════════════════════════════════════════════════════════════════
#  11. FIX NORMAL (WEB) PARAGRAPHS: change style to Normal for consistency
# ═══════════════════════════════════════════════════════════════════════════════

normal_style = doc.styles['Normal']
converted = 0
for p in doc.paragraphs:
    if p.style.name == 'Normal (Web)':
        p.style = normal_style
        converted += 1

print(f"OK 11: {converted} 'Normal (Web)' paragraphs converted to 'Normal'")

# ═══════════════════════════════════════════════════════════════════════════════
#  12. FIX BODY TEXT RUNS: ensure Times New Roman 12pt throughout
#      (Only fix runs that have no explicit size/font set, or wrong font)
# ═══════════════════════════════════════════════════════════════════════════════

fixed_runs = 0
for p in doc.paragraphs:
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3'):
        continue  # already handled
    for run in p.runs:
        changed = False
        if run.font.name not in (None, 'Times New Roman', 'Symbol'):
            run.font.name = 'Times New Roman'
            changed = True
        if run.font.size and run.font.size != Pt(12):
            # Only fix if not a deliberate different size (captions etc.)
            if run.font.size < Pt(8) or run.font.size > Pt(18):
                run.font.size = Pt(12)
                changed = True
        if changed:
            fixed_runs += 1

print(f"OK 12: Fixed font on {fixed_runs} body runs")

# ═══════════════════════════════════════════════════════════════════════════════
#  13. CHAPTER PAGE BREAKS: ensure each chapter H1 has a page break
#      (already handled via style, but also add explicit pageBreak paragraph
#       if a chapter follows directly after another without section break)
# ═══════════════════════════════════════════════════════════════════════════════

# The style-level pageBreakBefore handles this. Verify chapter headings all have it.
chapter_headings_fixed = 0
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and 'Chapter' in p.text:
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn('w:pageBreakBefore')) is None:
            pbr = OxmlElement('w:pageBreakBefore')
            pPr.append(pbr)
            chapter_headings_fixed += 1

print(f"OK 13: Page-break-before verified on chapter headings ({chapter_headings_fixed} fixed)")

# ═══════════════════════════════════════════════════════════════════════════════
#  14. FIX FIGURE/TABLE CAPTION PARAGRAPHS
#      Captions should be 11pt, centered, no first-line indent, italic optional
# ═══════════════════════════════════════════════════════════════════════════════

import re
caption_pattern = re.compile(r'^(Figure|Table|Fig\.)\s', re.IGNORECASE)

for p in doc.paragraphs:
    txt = p.text.strip()
    if caption_pattern.match(txt):
        pPr = p._p.get_or_add_pPr()
        set_jc(pPr, 'center')
        clear_first_line_indent(pPr)
        set_spacing(pPr, line='240', before='60', after='120')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

print("OK 14: Figure/table captions formatted (11pt, centered)")

# ═══════════════════════════════════════════════════════════════════════════════
#  15. ENSURE IMAGES / TABLES IN DOCUMENT ARE CENTERED
# ═══════════════════════════════════════════════════════════════════════════════

from docx.oxml.ns import qn as _qn
# Center paragraphs that contain only images (drawing/inline elements)
for p in doc.paragraphs:
    has_drawing = p._p.find('.//' + qn('w:drawing')) is not None
    has_pict    = p._p.find('.//' + qn('w:pict'))    is not None
    if has_drawing or has_pict:
        pPr = p._p.get_or_add_pPr()
        set_jc(pPr, 'center')
        clear_first_line_indent(pPr)

print("OK 15: Image-containing paragraphs centered")

# ═══════════════════════════════════════════════════════════════════════════════
#  16. SAVE
# ═══════════════════════════════════════════════════════════════════════════════

out_path = r'C:\Users\vvmar\Downloads\Graduation_Project_Formatted.docx'
doc.save(out_path)
print(f"\nSaved -> {out_path}")
print("Open the file in Word and press Ctrl+A then F9 to refresh the TOC.")
